from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path(__file__).resolve().parent
DOCX_PATH = OUT_DIR / "CXG-bt设备网管通信控制程序_bt_agent_ui_说明书.docx"

FONT = "Microsoft YaHei"
MONO = "Consolas"
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(35, 35, 35)
MUTED = RGBColor(95, 95, 95)
TABLE_HEADER = "E8EEF5"
CALLOUT = "F4F6F9"


def set_run_font(run, *, name=FONT, size=None, bold=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_paragraph_spacing(paragraph, *, before=0, after=6, line=1.25):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, *, bold=False, color=None, size=9.5, fill=None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    set_paragraph_spacing(paragraph, after=0, line=1.15)
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color or INK)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if fill:
        set_cell_shading(cell, fill)


def set_table_geometry(table, widths_in: list[float]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    total = sum(widths_in)
    tbl_w.set(qn("w:w"), str(int(total * 1440)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    for row_idx, row in enumerate(table.rows):
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:cantSplit")) is None:
            tr_pr.append(OxmlElement("w:cantSplit"))
        if row_idx == 0 and tr_pr.find(qn("w:tblHeader")) is None:
            tbl_header = OxmlElement("w:tblHeader")
            tbl_header.set(qn("w:val"), "true")
            tr_pr.append(tbl_header)
        for idx, width in enumerate(widths_in):
            if idx >= len(row.cells):
                continue
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")
            tc_mar = tc_pr.find(qn("w:tcMar"))
            if tc_mar is None:
                tc_mar = OxmlElement("w:tcMar")
                tc_pr.append(tc_mar)
            for side, value in (("top", "80"), ("bottom", "80"), ("start", "120"), ("end", "120")):
                el = tc_mar.find(qn(f"w:{side}"))
                if el is None:
                    el = OxmlElement(f"w:{side}")
                    tc_mar.append(el)
                el.set(qn("w:w"), value)
                el.set(qn("w:type"), "dxa")


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True, color=DARK_BLUE, fill=TABLE_HEADER)
    for row_values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_values):
            set_cell_text(cells[idx], value)
    set_table_geometry(table, widths)
    doc.add_paragraph()
    return table


def add_heading(doc: Document, text: str, level: int) -> None:
    doc.add_heading(text, level=level)


def add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    set_paragraph_spacing(p)
    run = p.add_run(text)
    set_run_font(run, size=11, color=INK)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)
        for run in p.runs:
            set_run_font(run, size=10.5)


def add_numbers(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)
        for run in p.runs:
            set_run_font(run, size=10.5)


def add_callout(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [6.5])
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, CALLOUT)
    paragraph = cell.paragraphs[0]
    set_paragraph_spacing(paragraph, after=2, line=1.2)
    run = paragraph.add_run(title)
    set_run_font(run, size=10.5, bold=True, color=DARK_BLUE)
    paragraph.add_run("\n")
    run = paragraph.add_run(body)
    set_run_font(run, size=10, color=INK)
    doc.add_paragraph()


def add_code_block(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=2, after=8, line=1.05)
    for line in text.rstrip().splitlines():
        run = p.add_run(line + "\n")
        set_run_font(run, name=MONO, size=8.5, color=RGBColor(30, 30, 30))


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, size=9, color=MUTED)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r = paragraph.add_run()._r
    r.append(fld_begin)
    r.append(instr)
    r.append(fld_end)
    run = paragraph.add_run(" 页")
    set_run_font(run, size=9, color=MUTED)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    header = section.header.paragraphs[0]
    header.text = ""
    run = header.add_run("bt_agent_ui 现场操作说明书")
    set_run_font(run, size=9, color=MUTED)
    footer = section.footer.paragraphs[0]
    add_page_number(footer)


def build_manual() -> None:
    doc = Document()
    configure_styles(doc)
    configure_document(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(title, after=3, line=1.15)
    run = title.add_run("CXG-bt设备网管通信控制程序")
    set_run_font(run, size=24, bold=True, color=DARK_BLUE)
    subtitle = doc.add_paragraph()
    set_paragraph_spacing(subtitle, after=12, line=1.15)
    run = subtitle.add_run("bt_agent_ui 详细说明书")
    set_run_font(run, size=15, color=BLUE)

    add_table(
        doc,
        ["项目", "内容"],
        [
            ["适用对象", "现场运维人员、交付人员、维护开发人员"],
            ["程序名称", "CXG-bt设备网管通信控制程序（bt_agent_ui）"],
            ["核心用途", "配置、启动、停止和监控 bt_agent 子进程，查看 BT UDP 通信质量与磁盘告警状态"],
            ["生成日期", date.today().isoformat()],
            ["依据文件", "bt_agent/bt_agent_ui.py、bt_agent/bt_agent.py、deploy/windows_host/README.md、BT_NMS运维手册"],
        ],
        [1.35, 5.15],
    )

    add_callout(
        doc,
        "阅读建议",
        "现场日常使用优先看第 3 至第 6 章；出现异常时看第 7 至第 9 章；部署和文件位置看第 2 章与第 10 章。",
    )

    add_heading(doc, "1. 程序定位", 1)
    add_body(
        doc,
        "bt_agent_ui 是 BT 采集程序的桌面控制界面。它本身不直接处理所有 UDP 报文，而是负责维护配置、启动 bt_agent 子进程、读取子进程输出的状态 JSON，并把运行状态、设备 IP 通信质量、日志和磁盘空间告警展示给现场人员。",
    )
    add_bullets(
        doc,
        [
            "适用场景：Windows 宿主机直接运行 BT agent，不通过 Docker 运行采集进程。",
            "数据方向：BT 设备通过 UDP 向本机监听端口发送报文，bt_agent 校验报文后写入 Redis Stream。",
            "命令方向：bt_agent 从 Redis 命令 Stream 读取下发命令，再通过 UDP 发回目标设备 IP。",
            "界面职责：统一配置 Redis、UDP、Stream、屏蔽 IP、磁盘告警与通信质量规则。",
            "保护部署：现场推荐使用编译后的 exe 包，通过 scripts\\run_bt_agent_ui.bat 启动。",
        ],
    )

    add_heading(doc, "2. 启动入口与运行文件", 1)
    add_table(
        doc,
        ["场景", "启动方式", "说明"],
        [
            ["现场保护部署", "scripts\\run_bt_agent_ui.bat", "从 windows_agents 包启动 bt_agent_ui.exe。目标机器不需要安装 Python。"],
            ["源码维护环境", "python bt_agent\\bt_agent_ui.py", "仅用于内部维护或调试，需要 Python、PySide6、redis 等依赖。"],
            ["批量启动", "scripts\\run_agents.bat", "同时打开 bt_agent_ui 和 sy_agent_ui；脚本窗口用于现场观察脚本输出。"],
        ],
        [1.35, 2.35, 2.8],
    )
    add_body(
        doc,
        "保护部署时必须复制整个 windows_agents 目录，不要只复制单个 exe。Nuitka standalone 程序依赖同目录 DLL、资源文件和 apps 目录结构。",
    )
    add_table(
        doc,
        ["文件", "位置", "用途"],
        [
            ["config.json", "%ProgramData%\\BT_NMS\\bt_agent\\config.json", "UI 保存的主配置文件，独立启动 bt_agent.exe 时也会读取。"],
            ["runtime_config.json", "%ProgramData%\\BT_NMS\\bt_agent\\runtime_config.json", "UI 点击启动时写出的运行配置，通过 BT_AGENT_CONFIG_JSON 传给 bt_agent。"],
            ["bt_agent_ui.sqlite3", "%ProgramData%\\BT_NMS\\bt_agent\\bt_agent_ui.sqlite3", "保存 UI 状态、窗口位置、自动启动、锁定状态、磁盘告警和通信质量规则。"],
            ["诊断包", "用户选择的 zip 输出路径", "导出 summary.json、config.json、runtime_config.json、recent.log 供排障。"],
        ],
        [1.35, 2.85, 2.3],
    )

    add_heading(doc, "3. 主界面结构", 1)
    add_table(
        doc,
        ["区域", "包含内容", "使用要点"],
        [
            ["控制", "启动/停止、保存设置、导入导出、诊断包、告警声、自动启动、锁定", "日常操作入口。启动前会自动保存配置并生成 runtime_config.json。"],
            ["通信质量概览", "主机状态指标和设备 IP 表", "用于判断进程、UDP、Redis、命令线程、设备在线和报文质量。"],
            ["配置与监控", "配置JSON、基础参数、屏蔽IP、磁盘告警、通信质量规则", "修改配置后点击保存设置；正在运行的子进程需要重启后读取新配置。"],
            ["日志", "最近可见日志", "非状态行会显示在日志框；状态行以 [BT_STATUS] 开头并被 UI 解析成指标。"],
        ],
        [1.2, 2.7, 2.6],
    )

    add_heading(doc, "4. 顶部按钮与状态栏", 1)
    add_table(
        doc,
        ["控件", "作用", "现场注意"],
        [
            ["启动 / 停止", "启动或终止 bt_agent 子进程。运行中按钮显示“停止”。", "启动失败通常与配置 JSON、端口占用、Redis 依赖或程序路径有关。"],
            ["保存设置", "校验配置 JSON，保存到 config.json 与 sqlite。", "基础参数页的改动会同步到配置 JSON；JSON 手工改错会保存失败。"],
            ["导入 JSON / 导出 JSON", "导入或备份运行配置。", "保护部署现场建议只使用 JSON 配置。"],
            ["导入 .py(开发) / 导出 .py(开发)", "兼容旧 CONFIG = {...} Python 配置。", "仅内部维护使用，现场不推荐。"],
            ["导出诊断包", "打包当前状态、配置、运行配置和最近日志。", "故障上报前优先导出，避免现场信息丢失。"],
            ["磁盘告警声音", "启用或禁用磁盘空间告警声音。", "关闭后仍会显示磁盘告警状态，但不播放声音。"],
            ["暂停告警声", "当前磁盘告警期间临时静音。", "磁盘告警解除后会自动恢复下一次告警能力。"],
            ["试音", "播放一次磁盘告警声音。", "用于确认现场音箱、声卡和音量。"],
            ["自动启动", "下次打开 UI 后自动启动 bt_agent。", "适合无人值守现场，变更配置前建议先停止子进程。"],
            ["锁定 / 解锁", "锁定后禁用配置和启动按钮，关闭窗口也会被拦截。", "解锁密码为 whbt；解锁状态下点击关闭会同步停止 bt_agent 子进程并退出 UI。"],
        ],
        [1.45, 2.55, 2.5],
    )
    add_table(
        doc,
        ["状态", "含义", "异常判断"],
        [
            ["进程", "bt_agent 子进程当前运行状态。", "已停止、异常退出、停止中都需要结合日志判断。"],
            ["通信质量", "综合进程、UDP、Redis、命令线程、发送错误和 IP 在线状态。", "出现“异常”时优先看最近问题和日志。"],
            ["UDP", "本机 UDP 监听状态。", "运行中但未监听通常是端口占用或绑定地址错误。"],
            ["Redis", "bt_agent 到 Redis 的连接状态。", "断开会影响报文写入和命令处理。"],
            ["磁盘告警", "C/D 盘或开发环境映射盘的空间告警状态。", "告警中、已静音、试音中、正常。"],
            ["最近问题", "当前最主要的一条异常。", "例如 Redis断开、UDP监听异常、存在离线IP。"],
            ["运行配置", "UI 启动子进程时写出的 runtime_config.json 路径。", "排障时可打开该文件确认子进程实际读取的配置。"],
        ],
        [1.2, 2.6, 2.7],
    )

    add_heading(doc, "5. 日常操作流程", 1)
    add_heading(doc, "5.1 启动前检查", 2)
    add_numbers(
        doc,
        [
            "确认 Redis 服务已启动，BT 默认连接 127.0.0.1:36379；如果后端在另一台机器，Redis 主机应改为后端机器 IP。",
            "确认 BT 设备报文目标 IP 指向本机，目标端口与 UI 基础参数中的 UDP 监听端口一致，默认 38315。",
            "确认 Windows 防火墙放行 UDP 监听端口和 Redis 端口访问。",
            "如需过滤测试设备或异常设备，先在“屏蔽IP”页加入源 IP。",
            "点击“保存设置”，确认没有 JSON 格式错误。",
        ],
    )
    add_heading(doc, "5.2 启动与停止", 2)
    add_numbers(
        doc,
        [
            "点击“启动”。UI 会保存当前设置、写出 runtime_config.json，并设置 BT_AGENT_CONFIG_JSON 后启动 bt_agent。",
            "启动后按钮变为“停止”，进程状态变为“运行中”，日志出现 started bt_agent pid=...。",
            "观察 UDP、Redis、通信质量三项。如果 Redis 未连接，先处理 Redis；如果 UDP 异常，先查端口占用和监听地址。",
            "需要停机维护时点击“停止”。正常手动停止后进程显示“已停止（手动）”。",
            "如果 UI 处于解锁状态，也可以直接点击窗口关闭按钮；程序会先停止 bt_agent 子进程，再保存 UI 状态并退出。",
        ],
    )
    add_heading(doc, "5.3 修改配置并生效", 2)
    add_numbers(
        doc,
        [
            "停止正在运行的 bt_agent 子进程，避免现场变更期间状态混乱。",
            "在基础参数、屏蔽 IP、磁盘告警或通信质量规则页修改参数。",
            "点击“保存设置”，确认日志显示 local settings saved。",
            "重新点击“启动”，让子进程读取新的 runtime_config.json。",
        ],
    )

    add_heading(doc, "6. 配置页说明", 1)
    add_heading(doc, "6.1 基础参数", 2)
    add_table(
        doc,
        ["字段", "默认值 / 范围", "说明"],
        [
            ["UDP 监听地址", "0.0.0.0", "监听所有网卡。固定网卡调试时可填本机指定 IP。"],
            ["UDP 监听端口", "38315 / 1-65535", "BT 设备发送 UDP 报文的目标端口。"],
            ["Redis 主机", "127.0.0.1", "Redis 服务地址。Docker 后端不在本机时应填写后端机器 IP。"],
            ["Redis 端口", "36379 / 1-65535", "BT 系统默认 Redis 端口。"],
            ["Packet Stream", "stream:udp:packets", "有效报文写入的 Redis Stream key。"],
            ["CMD Stream", "stream:udp:cmd", "控制命令读取的 Redis Stream key。"],
            ["CMD Group", "udp-agent-cmd", "Redis Stream 消费组名。"],
            ["CMD Consumer", "udp-agent-cmd-0", "消费组内消费者名，多个实例需区分。"],
            ["启动重试(秒)", "2 / 1-600", "启动阶段 Redis 不可用时的重试间隔。"],
            ["阻塞毫秒", "2000 / 100-60000", "从命令 Stream 读取命令的阻塞等待时长。"],
            ["读取条数", "100 / 1-10000", "单次读取命令的最大条数。"],
            ["Packet Maxlen", "200000 / 1000-10000000", "Packet Stream 近似最大长度。"],
            ["CMD Maxlen", "50000 / 1000-10000000", "命令或应答 Stream 近似最大长度。"],
        ],
        [1.65, 1.7, 3.15],
    )

    add_heading(doc, "6.2 配置 JSON", 2)
    add_body(
        doc,
        "配置 JSON 页是高级编辑入口，保存时会整体解析 JSON 并套用默认模板做归一化。基础参数页修改后会自动回写到 JSON 编辑器。现场只需要改常用参数时，优先使用基础参数页和屏蔽 IP 页，不建议直接手写 JSON。",
    )
    add_code_block(
        doc,
        """{
  "udp": {"host": "0.0.0.0", "port": 38315},
  "redis": {
    "host": "127.0.0.1",
    "port": 36379,
    "packet_stream_key": "stream:udp:packets",
    "cmd_stream_key": "stream:udp:cmd",
    "cmd_group": "udp-agent-cmd",
    "cmd_consumer": "udp-agent-cmd-0",
    "startup_retry_sec": 2.0
  },
  "stream": {
    "block_ms": 2000,
    "count": 100,
    "packet_maxlen": 200000,
    "cmd_maxlen": 50000
  },
  "filters": {"blocked_ips": []}
}""",
    )

    add_heading(doc, "6.3 屏蔽 IP", 2)
    add_body(
        doc,
        "屏蔽 IP 列表用于丢弃指定源 IP 的 UDP 报文。命中的报文不会进入通信质量设备表，但会计入主机状态中的“屏蔽包”。",
    )
    add_table(
        doc,
        ["操作", "说明"],
        [
            ["新增", "打开 IPv4 分段输入框，输入一个要屏蔽的源 IP。"],
            ["删除选中", "删除列表中选中的 IP。"],
            ["批量粘贴", "每行输入一个 IPv4 地址，程序会自动去重并过滤非法值。"],
            ["保存生效", "点击保存设置；如果子进程已运行，建议重启后让过滤列表按新配置生效。"],
        ],
        [1.35, 5.15],
    )

    add_heading(doc, "6.4 磁盘告警", 2)
    add_body(
        doc,
        "磁盘告警用于监控 Windows C 盘和 D 盘剩余空间。默认不启用，阈值默认 10%，可选 5%、10%、15%、20%、25%、30%。程序每 30 秒检查一次磁盘空间。",
    )
    add_bullets(
        doc,
        [
            "剩余空间百分比小于等于阈值时进入告警。",
            "磁盘告警声音启用时会播放 disk_space_alarm.wav。",
            "暂停告警声只对当前告警周期有效；磁盘恢复正常后静音状态会清除。",
            "开发环境下 C/D 显示会映射为系统盘和用户盘；现场 Windows 显示为 C盘、D盘。",
        ],
    )

    add_heading(doc, "6.5 通信质量规则", 2)
    add_table(
        doc,
        ["规则", "默认值", "状态含义"],
        [
            ["注意阈值", "10 秒", "某个 IP 最近有效收包间隔超过该值后，状态变为“注意”。"],
            ["离线阈值", "30 秒", "某个 IP 最近有效收包间隔超过该值后，状态变为“离线”。必须大于注意阈值。"],
            ["近期坏帧/校验错", "按离线阈值判断", "如果只有坏帧或校验错且没有有效包，近期会显示“异常”，超时后显示“离线”。"],
        ],
        [1.4, 1.2, 3.9],
    )

    add_heading(doc, "7. 通信质量概览解释", 1)
    add_table(
        doc,
        ["主机指标", "含义"],
        [
            ["运行时长", "本次 bt_agent 子进程启动后的连续运行时间。"],
            ["发送队列", "待发送给设备的 UDP 命令队列长度；大于 0 表示存在积压。"],
            ["有效包", "帧头、帧尾和校验和均正确的 UDP 报文数量。"],
            ["坏帧", "帧头或帧尾不正确的报文数量。"],
            ["校验错", "校验和不匹配的报文数量。"],
            ["屏蔽包", "源 IP 命中屏蔽列表后被丢弃的报文数量。"],
            ["模拟量包", "功能码识别为模拟量的数据包数量。"],
            ["命令接收 / 命令应答", "从 Redis 命令 Stream 读取到命令及完成应答的数量。"],
            ["发送成功 / 发送失败", "向设备发送 UDP 命令的结果统计。"],
            ["Redis发布失败", "写入 Packet Stream 失败的次数。"],
            ["最近收包 / 最近发送", "最近一次有效收包或发包时间。"],
        ],
        [1.75, 4.75],
    )
    add_table(
        doc,
        ["IP 表列", "说明"],
        [
            ["IP", "设备源 IP 地址，按收到过报文的来源自动发现。"],
            ["状态", "正常、注意、离线或异常。"],
            ["最近收包", "距离该 IP 最近一次收到任意报文的时间。"],
            ["速率", "最近 10 秒内有效报文平均速率。"],
            ["有效包 / 坏帧 / 校验错", "按 IP 统计的报文质量计数。"],
            ["发送成功 / 发送失败", "向该 IP 下发命令的结果计数。"],
        ],
        [1.75, 4.75],
    )
    add_callout(
        doc,
        "综合通信质量判断",
        "进程未运行、UDP监听异常、Redis断开、命令线程异常、存在发送失败或存在离线 IP 时显示“异常”；有发送队列积压或注意 IP 时显示“注意”；其他情况显示“正常”。",
    )

    add_heading(doc, "8. 日志与诊断包", 1)
    add_body(
        doc,
        "日志区显示最近的可见日志，主要用于现场快速判断启动、停止、Redis、UDP 收包和发送错误。bt_agent 每秒输出一行以 [BT_STATUS] 开头的状态 JSON，UI 会解析这类状态行并刷新概览，不把它作为普通日志长时间堆在日志框里。",
    )
    add_table(
        doc,
        ["诊断包文件", "内容"],
        [
            ["summary.json", "导出时间、进程/UDP/Redis/通信质量/磁盘告警状态、主机指标、IP 状态、磁盘告警配置、通信质量规则。"],
            ["config.json", "UI 当前保存的本地配置。"],
            ["runtime_config.json", "如果存在，包含最近一次 UI 启动子进程时写出的运行配置。"],
            ["recent.log", "UI 内部保留的最近日志。"],
        ],
        [1.75, 4.75],
    )
    add_numbers(
        doc,
        [
            "出现现场问题后不要先重启覆盖现场状态，优先点击“导出诊断包”。",
            "把 zip 文件连同故障发生时间、现场设备 IP、是否刚改过配置一起交给维护人员。",
            "如果 UI 已锁定，需要先解锁后才能导出诊断包。",
        ],
    )

    add_heading(doc, "9. 常见故障处理", 1)
    add_table(
        doc,
        ["现象", "优先检查", "处理建议"],
        [
            ["UI 无法重复打开", "是否已有一个 bt_agent_ui 实例运行。", "程序有单实例锁；关闭已有窗口或确认前台 Python/bt_agent_ui 进程。"],
            ["启动失败", "配置 JSON 是否有效、bt_agent.exe 或 bt_agent.py 是否存在。", "先点击保存设置；保护部署确认整个 windows_agents 包完整。"],
            ["UDP 显示异常", "监听地址、端口、端口占用、防火墙。", "默认端口 38315；确认设备目标端口一致并放行 UDP。"],
            ["Redis 显示断开", "Redis 主机、端口、网络、防火墙、Docker 后端。", "本机默认 127.0.0.1:36379；跨机器部署时填写后端 IP。"],
            ["通信质量异常", "最近问题、主机状态、IP 表状态、发送失败。", "先区分是进程/Redis/UDP 故障，还是单个设备离线或坏帧。"],
            ["某个 IP 异常", "有效包、坏帧、校验错、最近收包。", "坏帧或校验错多时检查协议帧头帧尾、校验和和设备固件。"],
            ["设备不显示", "设备是否发包、是否被屏蔽、UDP 目标是否正确。", "查看屏蔽包计数，必要时清空屏蔽 IP 后重启。"],
            ["磁盘告警响铃", "C/D 盘剩余空间、阈值、是否需要临时静音。", "清理磁盘或调高阈值；只想临时停止声音用“暂停告警声”。"],
            ["窗口无法关闭", "是否处于锁定状态。", "点击“解锁”，输入 whbt 后再关闭。"],
            ["子进程异常退出后又启动", "是否为非手动退出。", "UI 会在异常退出约 3 秒后自动重启 bt_agent；持续重启时查看日志和诊断包。"],
        ],
        [1.35, 2.25, 2.9],
    )

    add_heading(doc, "10. 现场部署与备份建议", 1)
    add_bullets(
        doc,
        [
            "保护部署产物位于 deploy/windows_host/artifacts/windows_agents/，现场转移时压缩整个目录。",
            "现场启动入口优先使用 scripts\\run_bt_agent_ui.bat。",
            "运行数据集中在 %ProgramData%\\BT_NMS\\bt_agent\\，排障时重点关注 config.json、runtime_config.json 和 bt_agent_ui.sqlite3。",
            "每次大范围修改 Redis、UDP 或屏蔽 IP 前，先导出 JSON 配置并保存一份诊断包。",
            "现场长期运行建议启用“自动启动”，并在确认配置稳定后使用“锁定”防止误操作。",
        ],
    )
    add_code_block(
        doc,
        """推荐备份内容：
%ProgramData%\\BT_NMS\\bt_agent\\config.json
%ProgramData%\\BT_NMS\\bt_agent\\runtime_config.json
%ProgramData%\\BT_NMS\\bt_agent\\bt_agent_ui.sqlite3
导出的 bt_agent_ui_diag_YYYYMMDD_HHMMSS.zip""",
    )

    add_heading(doc, "11. 日常巡检清单", 1)
    add_table(
        doc,
        ["检查项", "正常标准", "异常动作"],
        [
            ["进程", "运行中", "启动或查看异常退出日志。"],
            ["UDP", "监听中", "检查端口占用、防火墙和设备目标端口。"],
            ["Redis", "正常", "检查 Redis 服务、host、port 和网络连通性。"],
            ["通信质量", "正常", "查看最近问题和 IP 表。"],
            ["发送队列", "0 或短时间自动回落", "持续大于 0 时检查 Redis 命令量和设备网络。"],
            ["坏帧/校验错", "长期不增长或低频偶发", "持续增长时检查协议帧和设备固件。"],
            ["离线 IP", "无关键设备离线", "检查设备供电、网络和目标端口。"],
            ["磁盘告警", "正常", "清理磁盘或调整阈值。"],
            ["日志", "无连续报错", "导出诊断包并记录时间点。"],
        ],
        [1.6, 2.45, 2.45],
    )

    add_heading(doc, "12. 附录：关键默认值", 1)
    add_table(
        doc,
        ["类别", "键", "默认值"],
        [
            ["UDP", "udp.host", "0.0.0.0"],
            ["UDP", "udp.port", "38315"],
            ["Redis", "redis.host", "127.0.0.1"],
            ["Redis", "redis.port", "36379"],
            ["Redis", "redis.packet_stream_key", "stream:udp:packets"],
            ["Redis", "redis.cmd_stream_key", "stream:udp:cmd"],
            ["Redis", "redis.cmd_group", "udp-agent-cmd"],
            ["Redis", "redis.cmd_consumer", "udp-agent-cmd-0"],
            ["Redis", "redis.startup_retry_sec", "2.0"],
            ["Stream", "stream.block_ms", "2000"],
            ["Stream", "stream.count", "100"],
            ["Stream", "stream.packet_maxlen", "200000"],
            ["Stream", "stream.cmd_maxlen", "50000"],
            ["过滤", "filters.blocked_ips", "[]"],
            ["磁盘告警", "threshold_percent", "10，默认 C/D 均未启用"],
            ["通信质量", "warn_after_sec / offline_after_sec", "10 / 30"],
            ["解锁密码", "SETTINGS_LOCK_PASSWORD", "whbt"],
        ],
        [1.25, 2.65, 2.6],
    )

    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build_manual()
